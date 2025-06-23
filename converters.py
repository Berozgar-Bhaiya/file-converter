import os
import logging
import subprocess
import tempfile
from abc import ABC, abstractmethod
from PIL import Image
import fitz  # PyMuPDF
import pdfkit
from docx import Document

class BaseConverter(ABC):
    """Base class for all file converters"""
    
    @property
    @abstractmethod
    def output_mimetype(self):
        """MIME type of the output file"""
        pass
    
    @abstractmethod
    def convert(self, input_path, output_dir):
        """Convert file and return path to output file"""
        pass

class PDFToPNGConverter(BaseConverter):
    """Convert PDF to PNG using PyMuPDF"""
    
    @property
    def output_mimetype(self):
        return 'application/zip'  # Return as ZIP for multiple pages
    
    def convert(self, input_path, output_dir):
        try:
            import zipfile
            
            # Open PDF
            pdf_document = fitz.open(input_path)
            
            # Generate base name
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            
            # Create a temporary directory for PNG files
            png_dir = os.path.join(output_dir, f"{base_name}_pages")
            os.makedirs(png_dir, exist_ok=True)
            
            png_files = []
            
            # Convert all pages to PNG
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                mat = fitz.Matrix(2.0, 2.0)  # 2x scaling for better quality
                pix = page.get_pixmap(matrix=mat)
                
                # Generate output filename for each page
                page_filename = f"{base_name}_page_{page_num + 1:03d}.png"
                page_output_path = os.path.join(png_dir, page_filename)
                
                # Save PNG
                pix.save(page_output_path)
                png_files.append(page_output_path)
                
                logging.info(f"Converted page {page_num + 1} to PNG: {page_output_path}")
            
            pdf_document.close()
            
            # Create ZIP file containing all PNG files
            zip_output_path = os.path.join(output_dir, f"{base_name}_pages.zip")
            
            with zipfile.ZipFile(zip_output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for png_file in png_files:
                    # Add file to ZIP with just the filename (not full path)
                    zipf.write(png_file, os.path.basename(png_file))
            
            # Clean up temporary PNG files
            import shutil
            shutil.rmtree(png_dir)
            
            logging.info(f"Successfully converted PDF to PNG files (ZIP): {zip_output_path}")
            logging.info(f"Total pages converted: {len(png_files)}")
            
            return zip_output_path
            
        except Exception as e:
            logging.error(f"PDF to PNG conversion failed: {str(e)}")
            raise Exception(f"PDF conversion failed: {str(e)}")

class PNGToJPGConverter(BaseConverter):
    """Convert PNG to JPG using Pillow"""
    
    @property
    def output_mimetype(self):
        return 'image/jpeg'
    
    def convert(self, input_path, output_dir):
        try:
            # Open PNG image
            with Image.open(input_path) as img:
                # Convert RGBA to RGB if necessary
                if img.mode in ('RGBA', 'LA', 'P'):
                    # Create white background
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'P':
                        img = img.convert('RGBA')
                    background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
                    img = background
                elif img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # Generate output filename
                base_name = os.path.splitext(os.path.basename(input_path))[0]
                output_path = os.path.join(output_dir, f"{base_name}.jpg")
                
                # Save as JPG with high quality
                img.save(output_path, 'JPEG', quality=95, optimize=True)
                
                logging.info(f"Successfully converted PNG to JPG: {output_path}")
                return output_path
                
        except Exception as e:
            logging.error(f"PNG to JPG conversion failed: {str(e)}")
            raise Exception(f"Image conversion failed: {str(e)}")

class MP4ToMP3Converter(BaseConverter):
    """Convert MP4 to MP3 using FFmpeg"""
    
    @property
    def output_mimetype(self):
        return 'audio/mpeg'
    
    def convert(self, input_path, output_dir):
        try:
            # Generate output filename
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}.mp3")
            
            # FFmpeg command
            cmd = [
                'ffmpeg',
                '-i', input_path,
                '-vn',  # Disable video
                '-acodec', 'mp3',
                '-ab', '192k',  # Audio bitrate
                '-ar', '44100',  # Sample rate
                '-y',  # Overwrite output file
                output_path
            ]
            
            # Run FFmpeg
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=300  # 5 minute timeout
            )
            
            if result.returncode != 0:
                logging.error(f"FFmpeg error: {result.stderr}")
                raise Exception(f"Audio conversion failed: {result.stderr}")
            
            if not os.path.exists(output_path):
                raise Exception("Converted file was not created")
            
            logging.info(f"Successfully converted MP4 to MP3: {output_path}")
            return output_path
            
        except subprocess.TimeoutExpired:
            logging.error("MP4 to MP3 conversion timed out")
            raise Exception("Conversion timed out. File may be too large.")
        except Exception as e:
            logging.error(f"MP4 to MP3 conversion failed: {str(e)}")
            raise Exception(f"Audio conversion failed: {str(e)}")

class DOCXToPDFConverter(BaseConverter):
    """Convert DOCX to PDF using pdfkit (requires wkhtmltopdf)"""
    
    @property
    def output_mimetype(self):
        return 'application/pdf'
    
    def convert(self, input_path, output_dir):
        try:
            # Read DOCX content
            doc = Document(input_path)
            
            # Extract text content
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            
            content = '\n'.join(full_text)
            
            if not content.strip():
                raise Exception("Document appears to be empty")
            
            # Create HTML content
            content_paragraphs = '<p>' + content.replace('\n', '</p><p>') + '</p>'
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    body {{
                        font-family: Arial, sans-serif;
                        line-height: 1.6;
                        margin: 40px;
                        color: #333;
                    }}
                    p {{
                        margin: 12px 0;
                    }}
                </style>
            </head>
            <body>
                {content_paragraphs}
            </body>
            </html>
            """
            
            # Generate output filename
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}.pdf")
            
            # PDF generation options
            options = {
                'page-size': 'A4',
                'margin-top': '0.75in',
                'margin-right': '0.75in',
                'margin-bottom': '0.75in',
                'margin-left': '0.75in',
                'encoding': "UTF-8",
                'no-outline': None,
                'enable-local-file-access': None
            }
            
            # Convert HTML to PDF using reportlab as primary method
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.lib.units import inch
                
                doc_pdf = SimpleDocTemplate(output_path, pagesize=letter,
                                          rightMargin=0.75*inch, leftMargin=0.75*inch,
                                          topMargin=0.75*inch, bottomMargin=0.75*inch)
                styles = getSampleStyleSheet()
                story = []
                
                for line in content.split('\n'):
                    if line.strip():
                        story.append(Paragraph(line, styles['Normal']))
                        story.append(Spacer(1, 12))
                
                doc_pdf.build(story)
                
            except ImportError:
                # Fallback to pdfkit if reportlab not available
                try:
                    pdfkit.from_string(html_content, output_path, options=options)
                except OSError as e:
                    if 'wkhtmltopdf' in str(e):
                        raise Exception("Neither reportlab nor wkhtmltopdf available for PDF conversion. Please install one of them.")
                    else:
                        raise e
            
            if not os.path.exists(output_path):
                raise Exception("PDF file was not created")
            
            logging.info(f"Successfully converted DOCX to PDF: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"DOCX to PDF conversion failed: {str(e)}")
            raise Exception(f"Document conversion failed: {str(e)}")

class DOCToPDFConverter(BaseConverter):
    """Convert DOC to PDF using LibreOffice"""
    
    @property
    def output_mimetype(self):
        return 'application/pdf'
    
    def convert(self, input_path, output_dir):
        try:
            # Generate output filename
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}.pdf")
            
            # Try LibreOffice headless conversion first
            try:
                cmd = [
                    'libreoffice',
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', output_dir,
                    input_path
                ]
                
                result = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    timeout=60
                )
                
                if result.returncode == 0 and os.path.exists(output_path):
                    logging.info(f"Successfully converted DOC to PDF: {output_path}")
                    return output_path
                else:
                    raise Exception("LibreOffice conversion failed")
                    
            except (subprocess.TimeoutExpired, FileNotFoundError, Exception):
                # Fallback: Try to read with python-docx and convert to PDF
                try:
                    from docx import Document
                    doc = Document(input_path)
                    
                    # Extract text content
                    full_text = []
                    for paragraph in doc.paragraphs:
                        full_text.append(paragraph.text)
                    
                    content = '\n'.join(full_text)
                    
                    if not content.strip():
                        raise Exception("Document appears to be empty")
                    
                    # Create PDF using reportlab
                    from reportlab.lib.pagesizes import letter
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                    from reportlab.lib.styles import getSampleStyleSheet
                    from reportlab.lib.units import inch
                    
                    doc_pdf = SimpleDocTemplate(output_path, pagesize=letter,
                                              rightMargin=0.75*inch, leftMargin=0.75*inch,
                                              topMargin=0.75*inch, bottomMargin=0.75*inch)
                    styles = getSampleStyleSheet()
                    story = []
                    
                    for line in content.split('\n'):
                        if line.strip():
                            story.append(Paragraph(line, styles['Normal']))
                            story.append(Spacer(1, 12))
                    
                    doc_pdf.build(story)
                    
                except ImportError:
                    raise Exception("Neither LibreOffice nor required Python libraries available for DOC conversion.")
            
            if not os.path.exists(output_path):
                raise Exception("PDF file was not created")
            
            logging.info(f"Successfully converted DOC to PDF: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"DOC to PDF conversion failed: {str(e)}")
            raise Exception(f"Document conversion failed: {str(e)}")

class EPUBToPDFConverter(BaseConverter):
    """Convert EPUB to PDF using reportlab"""
    
    @property
    def output_mimetype(self):
        return 'application/pdf'
    
    def convert(self, input_path, output_dir):
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            from html import unescape
            import re
            
            # Generate output filename
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}.pdf")
            
            # Extract text from EPUB
            content_text = []
            
            with zipfile.ZipFile(input_path, 'r') as epub_zip:
                # Find content files
                try:
                    # Try to read content.opf or similar manifest
                    manifest_files = [f for f in epub_zip.namelist() if f.endswith('.opf')]
                    
                    if manifest_files:
                        manifest_content = epub_zip.read(manifest_files[0]).decode('utf-8')
                        # Simple extraction - find xhtml/html files
                        html_files = [f for f in epub_zip.namelist() if f.endswith(('.xhtml', '.html'))]
                    else:
                        html_files = [f for f in epub_zip.namelist() if f.endswith(('.xhtml', '.html'))]
                    
                    for html_file in html_files:
                        try:
                            html_content = epub_zip.read(html_file).decode('utf-8')
                            # Simple HTML tag removal
                            clean_text = re.sub(r'<[^>]+>', '', html_content)
                            clean_text = unescape(clean_text)
                            clean_text = re.sub(r'\s+', ' ', clean_text).strip()
                            if clean_text:
                                content_text.append(clean_text)
                        except:
                            continue
                            
                except Exception:
                    # Fallback: try to extract any text content
                    for file_info in epub_zip.filelist:
                        if file_info.filename.endswith(('.txt', '.xhtml', '.html')):
                            try:
                                file_content = epub_zip.read(file_info).decode('utf-8')
                                clean_text = re.sub(r'<[^>]+>', '', file_content)
                                clean_text = unescape(clean_text)
                                clean_text = re.sub(r'\s+', ' ', clean_text).strip()
                                if clean_text:
                                    content_text.append(clean_text)
                            except:
                                continue
            
            if not content_text:
                raise Exception("No readable content found in EPUB file")
            
            full_content = '\n\n'.join(content_text)
            
            # Create PDF using reportlab
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.lib.units import inch
                
                doc_pdf = SimpleDocTemplate(output_path, pagesize=letter,
                                          rightMargin=0.75*inch, leftMargin=0.75*inch,
                                          topMargin=0.75*inch, bottomMargin=0.75*inch)
                styles = getSampleStyleSheet()
                story = []
                
                # Split content into paragraphs
                paragraphs = full_content.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        story.append(Paragraph(para.strip(), styles['Normal']))
                        story.append(Spacer(1, 12))
                
                doc_pdf.build(story)
                
            except ImportError:
                raise Exception("reportlab not available for EPUB to PDF conversion.")
            
            if not os.path.exists(output_path):
                raise Exception("PDF file was not created")
            
            logging.info(f"Successfully converted EPUB to PDF: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"EPUB to PDF conversion failed: {str(e)}")
            raise Exception(f"EPUB conversion failed: {str(e)}")

class JPGToPDFConverter(BaseConverter):
    """Convert JPG to PDF using Pillow"""
    
    @property
    def output_mimetype(self):
        return 'application/pdf'
    
    def convert(self, input_path, output_dir):
        try:
            # Open image
            with Image.open(input_path) as img:
                # Convert to RGB if necessary
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # Generate output filename
                base_name = os.path.splitext(os.path.basename(input_path))[0]
                output_path = os.path.join(output_dir, f"{base_name}.pdf")
                
                # Save as PDF
                img.save(output_path, 'PDF', resolution=100.0)
                
                logging.info(f"Successfully converted JPG to PDF: {output_path}")
                return output_path
                
        except Exception as e:
            logging.error(f"JPG to PDF conversion failed: {str(e)}")
            raise Exception(f"Image conversion failed: {str(e)}")

class PDFToDOCXConverter(BaseConverter):
    """Convert PDF to DOCX using pdf2docx"""
    
    @property
    def output_mimetype(self):
        return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    
    def convert(self, input_path, output_dir):
        try:
            # Generate output filename
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}.docx")
            
            try:
                from pdf2docx import Converter
                
                # Convert PDF to DOCX
                cv = Converter(input_path)
                cv.convert(output_path, start=0, end=None)
                cv.close()
                
            except ImportError:
                # Fallback: Extract text and create basic DOCX
                import fitz  # PyMuPDF
                from docx import Document
                
                # Extract text from PDF
                pdf_document = fitz.open(input_path)
                full_text = []
                
                for page_num in range(pdf_document.page_count):
                    page = pdf_document[page_num]
                    text = page.get_text()
                    if text.strip():
                        full_text.append(text)
                
                pdf_document.close()
                
                if not full_text:
                    raise Exception("No text content found in PDF")
                
                # Create DOCX document
                doc = Document()
                for text_block in full_text:
                    paragraphs = text_block.split('\n')
                    for para in paragraphs:
                        if para.strip():
                            doc.add_paragraph(para.strip())
                
                doc.save(output_path)
            
            if not os.path.exists(output_path):
                raise Exception("DOCX file was not created")
            
            logging.info(f"Successfully converted PDF to DOCX: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"PDF to DOCX conversion failed: {str(e)}")
            raise Exception(f"PDF conversion failed: {str(e)}")

class XLSXToPDFConverter(BaseConverter):
    """Convert XLSX to PDF using openpyxl and reportlab"""
    
    @property
    def output_mimetype(self):
        return 'application/pdf'
    
    def convert(self, input_path, output_dir):
        try:
            # Generate output filename
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}.pdf")
            
            try:
                from openpyxl import load_workbook
                from reportlab.lib.pagesizes import letter, landscape
                from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.lib.units import inch
                from reportlab.lib import colors
                
                # Load workbook
                wb = load_workbook(input_path, data_only=True)
                
                # Create PDF
                doc_pdf = SimpleDocTemplate(output_path, pagesize=landscape(letter),
                                          rightMargin=0.5*inch, leftMargin=0.5*inch,
                                          topMargin=0.5*inch, bottomMargin=0.5*inch)
                styles = getSampleStyleSheet()
                story = []
                
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    
                    # Add sheet title
                    story.append(Paragraph(f"Sheet: {sheet_name}", styles['Heading1']))
                    story.append(Spacer(1, 12))
                    
                    # Get data from sheet
                    data = []
                    max_col = 0
                    for row in ws.iter_rows(values_only=True):
                        if any(cell is not None for cell in row):
                            row_data = [str(cell) if cell is not None else '' for cell in row]
                            data.append(row_data)
                            max_col = max(max_col, len(row_data))
                    
                    if data:
                        # Normalize row lengths
                        for row in data:
                            while len(row) < max_col:
                                row.append('')
                        
                        # Create table
                        table = Table(data)
                        table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('FONTSIZE', (0, 0), (-1, 0), 10),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                            ('GRID', (0, 0), (-1, -1), 1, colors.black),
                            ('FONTSIZE', (0, 1), (-1, -1), 8),
                        ]))
                        
                        story.append(table)
                        story.append(Spacer(1, 20))
                
                doc_pdf.build(story)
                
            except ImportError:
                raise Exception("openpyxl and reportlab required for XLSX to PDF conversion.")
            
            if not os.path.exists(output_path):
                raise Exception("PDF file was not created")
            
            logging.info(f"Successfully converted XLSX to PDF: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"XLSX to PDF conversion failed: {str(e)}")
            raise Exception(f"Excel conversion failed: {str(e)}")

class PPTXToPDFConverter(BaseConverter):
    """Convert PPTX to PDF using LibreOffice"""
    
    @property
    def output_mimetype(self):
        return 'application/pdf'
    
    def convert(self, input_path, output_dir):
        try:
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}.pdf")
            
            cmd = [
                'libreoffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', output_dir,
                input_path
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            
            if result.returncode != 0 or not os.path.exists(output_path):
                raise Exception("LibreOffice conversion failed")
            
            logging.info(f"Successfully converted PPTX to PDF: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"PPTX to PDF conversion failed: {str(e)}")
            raise Exception(f"Presentation conversion failed: {str(e)}")

class PDFToTXTConverter(BaseConverter):
    """Convert PDF to TXT using PyMuPDF"""
    
    @property
    def output_mimetype(self):
        return 'text/plain'
    
    def convert(self, input_path, output_dir):
        try:
            import fitz
            
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}.txt")
            
            pdf_document = fitz.open(input_path)
            full_text = []
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                text = page.get_text()
                if text.strip():
                    full_text.append(text)
            
            pdf_document.close()
            
            if not full_text:
                raise Exception("No text content found in PDF")
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n\n'.join(full_text))
            
            logging.info(f"Successfully converted PDF to TXT: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"PDF to TXT conversion failed: {str(e)}")
            raise Exception(f"PDF text extraction failed: {str(e)}")

class XLSXToCSVConverter(BaseConverter):
    """Convert XLSX to CSV using openpyxl"""
    
    @property
    def output_mimetype(self):
        return 'text/csv'
    
    def convert(self, input_path, output_dir):
        try:
            from openpyxl import load_workbook
            import csv
            
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}.csv")
            
            wb = load_workbook(input_path, data_only=True)
            ws = wb.active  # Use the active sheet
            
            with open(output_path, 'w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)
                for row in ws.iter_rows(values_only=True):
                    writer.writerow([str(cell) if cell is not None else '' for cell in row])
            
            logging.info(f"Successfully converted XLSX to CSV: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"XLSX to CSV conversion failed: {str(e)}")
            raise Exception(f"Excel to CSV conversion failed: {str(e)}")

class CSVToXLSXConverter(BaseConverter):
    """Convert CSV to XLSX using openpyxl"""
    
    @property
    def output_mimetype(self):
        return 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    
    def convert(self, input_path, output_dir):
        try:
            from openpyxl import Workbook
            import csv
            
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}.xlsx")
            
            wb = Workbook()
            ws = wb.active
            
            with open(input_path, 'r', encoding='utf-8') as csvfile:
                reader = csv.reader(csvfile)
                for row_num, row in enumerate(reader, 1):
                    for col_num, value in enumerate(row, 1):
                        ws.cell(row=row_num, column=col_num, value=value)
            
            wb.save(output_path)
            
            logging.info(f"Successfully converted CSV to XLSX: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"CSV to XLSX conversion failed: {str(e)}")
            raise Exception(f"CSV to Excel conversion failed: {str(e)}")

class CSVToPDFConverter(BaseConverter):
    """Convert CSV to PDF using reportlab"""
    
    @property
    def output_mimetype(self):
        return 'application/pdf'
    
    def convert(self, input_path, output_dir):
        try:
            import csv
            from reportlab.lib.pagesizes import letter, landscape
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
            from reportlab.lib import colors
            
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}.pdf")
            
            # Read CSV data
            data = []
            with open(input_path, 'r', encoding='utf-8') as csvfile:
                reader = csv.reader(csvfile)
                for row in reader:
                    data.append(row)
            
            if not data:
                raise Exception("CSV file is empty")
            
            # Create PDF
            doc = SimpleDocTemplate(output_path, pagesize=landscape(letter))
            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            
            doc.build([table])
            
            logging.info(f"Successfully converted CSV to PDF: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"CSV to PDF conversion failed: {str(e)}")
            raise Exception(f"CSV conversion failed: {str(e)}")

class HTMLToPDFConverter(BaseConverter):
    """Convert HTML to PDF using pdfkit"""
    
    @property
    def output_mimetype(self):
        return 'application/pdf'
    
    def convert(self, input_path, output_dir):
        try:
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}.pdf")
            
            try:
                import pdfkit
                pdfkit.from_file(input_path, output_path)
            except:
                # Fallback using reportlab
                with open(input_path, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                
                # Simple HTML to text conversion for PDF
                import re
                from html import unescape
                clean_text = re.sub(r'<[^>]+>', '', html_content)
                clean_text = unescape(clean_text)
                
                from reportlab.lib.pagesizes import letter
                from reportlab.platypus import SimpleDocTemplate, Paragraph
                from reportlab.lib.styles import getSampleStyleSheet
                
                doc = SimpleDocTemplate(output_path, pagesize=letter)
                styles = getSampleStyleSheet()
                story = [Paragraph(clean_text, styles['Normal'])]
                doc.build(story)
            
            logging.info(f"Successfully converted HTML to PDF: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"HTML to PDF conversion failed: {str(e)}")
            raise Exception(f"HTML conversion failed: {str(e)}")

class PDFToHTMLConverter(BaseConverter):
    """Convert PDF to HTML using PyMuPDF"""
    
    @property
    def output_mimetype(self):
        return 'text/html'
    
    def convert(self, input_path, output_dir):
        try:
            import fitz
            
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}.html")
            
            pdf_document = fitz.open(input_path)
            html_content = ["<!DOCTYPE html><html><head><title>Converted PDF</title></head><body>"]
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                text = page.get_text()
                if text.strip():
                    html_content.append(f"<h2>Page {page_num + 1}</h2>")
                    paragraphs = text.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            html_content.append(f"<p>{para.strip()}</p>")
            
            html_content.append("</body></html>")
            pdf_document.close()
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(html_content))
            
            logging.info(f"Successfully converted PDF to HTML: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"PDF to HTML conversion failed: {str(e)}")
            raise Exception(f"PDF to HTML conversion failed: {str(e)}")

# Audio converters
class MP3ToWAVConverter(BaseConverter):
    """Convert MP3 to WAV using FFmpeg"""
    
    @property
    def output_mimetype(self):
        return 'audio/wav'
    
    def convert(self, input_path, output_dir):
        try:
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}.wav")
            
            cmd = ['ffmpeg', '-i', input_path, '-y', output_path]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
            
            if result.returncode != 0:
                raise Exception(f"FFmpeg error: {result.stderr}")
            
            logging.info(f"Successfully converted MP3 to WAV: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"MP3 to WAV conversion failed: {str(e)}")
            raise Exception(f"Audio conversion failed: {str(e)}")

class WAVToMP3Converter(BaseConverter):
    """Convert WAV to MP3 using FFmpeg"""
    
    @property
    def output_mimetype(self):
        return 'audio/mpeg'
    
    def convert(self, input_path, output_dir):
        try:
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}.mp3")
            
            cmd = ['ffmpeg', '-i', input_path, '-acodec', 'mp3', '-ab', '192k', '-y', output_path]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
            
            if result.returncode != 0:
                raise Exception(f"FFmpeg error: {result.stderr}")
            
            logging.info(f"Successfully converted WAV to MP3: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"WAV to MP3 conversion failed: {str(e)}")
            raise Exception(f"Audio conversion failed: {str(e)}")

# Video converters
class MP4ToAVIConverter(BaseConverter):
    """Convert MP4 to AVI using FFmpeg"""
    
    @property
    def output_mimetype(self):
        return 'video/x-msvideo'
    
    def convert(self, input_path, output_dir):
        try:
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}.avi")
            
            cmd = ['ffmpeg', '-i', input_path, '-c:v', 'libx264', '-c:a', 'aac', '-y', output_path]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=600)
            
            if result.returncode != 0:
                raise Exception(f"FFmpeg error: {result.stderr}")
            
            logging.info(f"Successfully converted MP4 to AVI: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"MP4 to AVI conversion failed: {str(e)}")
            raise Exception(f"Video conversion failed: {str(e)}")

class AVIToMP4Converter(BaseConverter):
    """Convert AVI to MP4 using FFmpeg"""
    
    @property
    def output_mimetype(self):
        return 'video/mp4'
    
    def convert(self, input_path, output_dir):
        try:
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}.mp4")
            
            cmd = ['ffmpeg', '-i', input_path, '-c:v', 'libx264', '-c:a', 'aac', '-y', output_path]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=600)
            
            if result.returncode != 0:
                raise Exception(f"FFmpeg error: {result.stderr}")
            
            logging.info(f"Successfully converted AVI to MP4: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"AVI to MP4 conversion failed: {str(e)}")
            raise Exception(f"Video conversion failed: {str(e)}")

# Image converters
class WEBPToJPGConverter(BaseConverter):
    """Convert WEBP to JPG using Pillow"""
    
    @property
    def output_mimetype(self):
        return 'image/jpeg'
    
    def convert(self, input_path, output_dir):
        try:
            with Image.open(input_path) as img:
                if img.mode in ('RGBA', 'LA', 'P'):
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'P':
                        img = img.convert('RGBA')
                    background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
                    img = background
                elif img.mode != 'RGB':
                    img = img.convert('RGB')
                
                base_name = os.path.splitext(os.path.basename(input_path))[0]
                output_path = os.path.join(output_dir, f"{base_name}.jpg")
                
                img.save(output_path, 'JPEG', quality=95)
                
                logging.info(f"Successfully converted WEBP to JPG: {output_path}")
                return output_path
                
        except Exception as e:
            logging.error(f"WEBP to JPG conversion failed: {str(e)}")
            raise Exception(f"Image conversion failed: {str(e)}")

class BMPToPNGConverter(BaseConverter):
    """Convert BMP to PNG using Pillow"""
    
    @property
    def output_mimetype(self):
        return 'image/png'
    
    def convert(self, input_path, output_dir):
        try:
            with Image.open(input_path) as img:
                base_name = os.path.splitext(os.path.basename(input_path))[0]
                output_path = os.path.join(output_dir, f"{base_name}.png")
                
                img.save(output_path, 'PNG')
                
                logging.info(f"Successfully converted BMP to PNG: {output_path}")
                return output_path
                
        except Exception as e:
            logging.error(f"BMP to PNG conversion failed: {str(e)}")
            raise Exception(f"Image conversion failed: {str(e)}")

class CompressPDFConverter(BaseConverter):
    """Compress PDF using PyMuPDF"""
    
    @property
    def output_mimetype(self):
        return 'application/pdf'
    
    def convert(self, input_path, output_dir):
        try:
            import fitz
            
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}_compressed.pdf")
            
            pdf_document = fitz.open(input_path)
            
            # Save with compression
            pdf_document.save(output_path, garbage=4, deflate=True, clean=True)
            pdf_document.close()
            
            logging.info(f"Successfully compressed PDF: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"PDF compression failed: {str(e)}")
            raise Exception(f"PDF compression failed: {str(e)}")

class JPGToWEBPConverter(BaseConverter):
    """Convert JPG to WEBP using Pillow"""
    
    @property
    def output_mimetype(self):
        return 'image/webp'
    
    def convert(self, input_path, output_dir):
        try:
            with Image.open(input_path) as img:
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                
                base_name = os.path.splitext(os.path.basename(input_path))[0]
                output_path = os.path.join(output_dir, f"{base_name}.webp")
                
                img.save(output_path, 'WEBP', quality=85, method=6)
                
                logging.info(f"Successfully converted JPG to WEBP: {output_path}")
                return output_path
                
        except Exception as e:
            logging.error(f"JPG to WEBP conversion failed: {str(e)}")
            raise Exception(f"Image conversion failed: {str(e)}")

class JPGToPNGConverter(BaseConverter):
    """Convert JPG to PNG using Pillow"""
    
    @property
    def output_mimetype(self):
        return 'image/png'
    
    def convert(self, input_path, output_dir):
        try:
            with Image.open(input_path) as img:
                base_name = os.path.splitext(os.path.basename(input_path))[0]
                output_path = os.path.join(output_dir, f"{base_name}.png")
                
                img.save(output_path, 'PNG')
                
                logging.info(f"Successfully converted JPG to PNG: {output_path}")
                return output_path
                
        except Exception as e:
            logging.error(f"JPG to PNG conversion failed: {str(e)}")
            raise Exception(f"Image conversion failed: {str(e)}")

class MKVToMP4Converter(BaseConverter):
    """Convert MKV to MP4 using FFmpeg"""
    
    @property
    def output_mimetype(self):
        return 'video/mp4'
    
    def convert(self, input_path, output_dir):
        try:
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}.mp4")
            
            cmd = ['ffmpeg', '-i', input_path, '-c:v', 'libx264', '-c:a', 'aac', '-y', output_path]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=600)
            
            if result.returncode != 0:
                raise Exception(f"FFmpeg error: {result.stderr}")
            
            logging.info(f"Successfully converted MKV to MP4: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"MKV to MP4 conversion failed: {str(e)}")
            raise Exception(f"Video conversion failed: {str(e)}")

class MP4ToWEBMConverter(BaseConverter):
    """Convert MP4 to WebM using FFmpeg"""
    
    @property
    def output_mimetype(self):
        return 'video/webm'
    
    def convert(self, input_path, output_dir):
        try:
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}.webm")
            
            cmd = ['ffmpeg', '-i', input_path, '-c:v', 'libvpx-vp9', '-c:a', 'libopus', '-y', output_path]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=600)
            
            if result.returncode != 0:
                raise Exception(f"FFmpeg error: {result.stderr}")
            
            logging.info(f"Successfully converted MP4 to WebM: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"MP4 to WebM conversion failed: {str(e)}")
            raise Exception(f"Video conversion failed: {str(e)}")

class OGGToMP3Converter(BaseConverter):
    """Convert OGG to MP3 using FFmpeg"""
    
    @property
    def output_mimetype(self):
        return 'audio/mpeg'
    
    def convert(self, input_path, output_dir):
        try:
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}.mp3")
            
            cmd = ['ffmpeg', '-i', input_path, '-acodec', 'mp3', '-ab', '192k', '-y', output_path]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
            
            if result.returncode != 0:
                raise Exception(f"FFmpeg error: {result.stderr}")
            
            logging.info(f"Successfully converted OGG to MP3: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"OGG to MP3 conversion failed: {str(e)}")
            raise Exception(f"Audio conversion failed: {str(e)}")

class MergePDFsConverter(BaseConverter):
    """Merge multiple PDF files into one using PyMuPDF"""
    
    @property
    def output_mimetype(self):
        return 'application/pdf'
    
    def convert(self, input_paths, output_dir):
        """
        Convert method for merging PDFs
        input_paths: list of PDF file paths to merge
        output_dir: directory to save the merged PDF
        """
        try:
            import fitz
            
            # Create new PDF document
            merged_pdf = fitz.open()
            total_pages = 0
            
            # Process each input PDF
            for i, pdf_path in enumerate(input_paths):
                try:
                    # Open source PDF
                    source_pdf = fitz.open(pdf_path)
                    page_count = source_pdf.page_count
                    
                    # Insert all pages from source PDF
                    merged_pdf.insert_pdf(source_pdf)
                    total_pages += page_count
                    
                    # Close source PDF after processing
                    source_pdf.close()
                    
                    logging.info(f"Added PDF {i+1}: {os.path.basename(pdf_path)} ({page_count} pages)")
                    
                except Exception as e:
                    logging.error(f"Failed to process PDF {pdf_path}: {str(e)}")
                    continue
            
            if merged_pdf.page_count == 0:
                merged_pdf.close()
                raise Exception("No valid PDF pages found to merge")
            
            # Generate output filename
            output_path = os.path.join(output_dir, "merged_document.pdf")
            
            # Save merged PDF
            merged_pdf.save(output_path)
            final_page_count = merged_pdf.page_count
            merged_pdf.close()
            
            logging.info(f"Successfully merged {len(input_paths)} PDFs into: {output_path}")
            logging.info(f"Total pages in merged PDF: {final_page_count}")
            
            return output_path
            
        except Exception as e:
            logging.error(f"PDF merge failed: {str(e)}")
            raise Exception(f"PDF merge failed: {str(e)}")

class MergeImagesConverter(BaseConverter):
    """Merge multiple images into one PDF using Pillow"""
    
    @property
    def output_mimetype(self):
        return 'application/pdf'
    
    def convert(self, input_paths, output_dir):
        """
        Convert method for merging images into PDF
        input_paths: list of image file paths to merge
        output_dir: directory to save the merged PDF
        """
        try:
            from PIL import Image
            
            images = []
            
            # Process each input image
            for i, img_path in enumerate(input_paths):
                try:
                    # Open and convert image
                    with Image.open(img_path) as img:
                        # Convert to RGB if necessary
                        if img.mode != 'RGB':
                            img = img.convert('RGB')
                        
                        # Resize image if too large (max width/height: 2000px)
                        max_size = 2000
                        if img.width > max_size or img.height > max_size:
                            img.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)
                        
                        # Store the processed image
                        images.append(img.copy())
                        
                    logging.info(f"Added image {i+1}: {os.path.basename(img_path)} ({img.width}x{img.height})")
                    
                except Exception as e:
                    logging.error(f"Failed to process image {img_path}: {str(e)}")
                    continue
            
            if not images:
                raise Exception("No valid images found to merge")
            
            # Generate output filename
            output_path = os.path.join(output_dir, "merged_images.pdf")
            
            # Save all images as a single PDF
            if len(images) == 1:
                images[0].save(output_path, 'PDF', resolution=100.0)
            else:
                # Save first image as PDF and append others
                images[0].save(
                    output_path, 
                    'PDF', 
                    resolution=100.0,
                    save_all=True,
                    append_images=images[1:]
                )
            
            logging.info(f"Successfully merged {len(images)} images into: {output_path}")
            
            return output_path
            
        except Exception as e:
            logging.error(f"Image merge failed: {str(e)}")
            raise Exception(f"Image merge failed: {str(e)}")

class TXTToPDFConverter(BaseConverter):
    """Convert TXT to PDF using reportlab"""
    
    @property
    def output_mimetype(self):
        return 'application/pdf'
    
    def convert(self, input_path, output_dir):
        try:
            # Read text file
            with open(input_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            if not content.strip():
                raise Exception("Text file appears to be empty")
            
            # Generate output filename
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}.pdf")
            
            # Create PDF using reportlab
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.lib.units import inch
                from reportlab.platypus import Preformatted
                
                doc_pdf = SimpleDocTemplate(output_path, pagesize=letter,
                                          rightMargin=0.75*inch, leftMargin=0.75*inch,
                                          topMargin=0.75*inch, bottomMargin=0.75*inch)
                styles = getSampleStyleSheet()
                story = []
                
                # Use preformatted text to preserve formatting
                preformatted_text = Preformatted(content, styles['Code'])
                story.append(preformatted_text)
                
                doc_pdf.build(story)
                
            except ImportError:
                raise Exception("reportlab not available for TXT to PDF conversion.")
            
            if not os.path.exists(output_path):
                raise Exception("PDF file was not created")
            
            logging.info(f"Successfully converted TXT to PDF: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"TXT to PDF conversion failed: {str(e)}")
            raise Exception(f"Text conversion failed: {str(e)}")
